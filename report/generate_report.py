"""
Generate a polished, academically styled DOCX report for the
LLM Context Compression as a 0/1 Knapsack Problem project.
Includes rendered PNG charts embedded directly in the document.
"""

import os
import csv
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── paths ────────────────────────────────────────────────────────────────────
BASE   = r"C:\Users\reuve\OneDrive\Documents\Knapsack_Optimization_Problem"
REPORT = os.path.join(BASE, "report")
OUT    = os.path.join(REPORT, "final_report.docx")

CSV_DP     = os.path.join(BASE, "results", "exact_dp",     "fresh_full_exact_dp",     "summary.csv")
CSV_GR     = os.path.join(BASE, "results", "greedy_ratio", "fresh_full_greedy_ratio", "summary.csv")
CSV_GF     = os.path.join(BASE, "results", "greedy_refine","fresh_full_greedy_refine","summary.csv")

# ── colour palette ────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1A, 0x23, 0x4E)
SLATE   = RGBColor(0x2E, 0x3A, 0x59)
TEAL    = RGBColor(0x0D, 0x7C, 0x8F)
GOLD    = RGBColor(0xC8, 0x9A, 0x2B)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT   = RGBColor(0xF0, 0xF4, 0xF8)
TH_BG   = RGBColor(0x1A, 0x23, 0x4E)
ALT_ROW = RGBColor(0xE8, 0xEE, 0xF5)

# ── chart colours ─────────────────────────────────────────────────────────────
C_DP  = "#1A234E"
C_GR  = "#0D7C8F"
C_GF  = "#C89A2B"

# ─────────────────────────────────────────────────────────────────────────────
# CSV helpers
# ─────────────────────────────────────────────────────────────────────────────

def read_csv(path):
    rows = []
    with open(path, newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    return rows


def get_runtime_by_merge(rows, budget):
    """Return {merge_size: avg_runtime_sec} for a given budget."""
    result = {}
    for r in rows:
        if int(r["budget_tokens"]) == budget:
            result[int(r["merge_size"])] = float(r["avg_runtime_sec"])
    return result


# ─────────────────────────────────────────────────────────────────────────────
# Chart generation
# ─────────────────────────────────────────────────────────────────────────────

def make_runtime_chart(dp_rows, gr_rows, gf_rows, budget, out_path):
    merge_sizes = [10, 20, 30, 40, 50]

    dp_rt = [get_runtime_by_merge(dp_rows, budget).get(m, 0) for m in merge_sizes]
    gr_rt = [get_runtime_by_merge(gr_rows, budget).get(m, 0) for m in merge_sizes]
    gf_rt = [get_runtime_by_merge(gf_rows, budget).get(m, 0) for m in merge_sizes]

    fig, ax = plt.subplots(figsize=(7, 4))
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    ax.plot(merge_sizes, dp_rt, marker="o", linewidth=2.2, color=C_DP,
            label="Exact-DP", zorder=3)
    ax.plot(merge_sizes, gr_rt, marker="s", linewidth=2.2, color=C_GR,
            label="Greedy-Ratio", zorder=3)
    ax.plot(merge_sizes, gf_rt, marker="^", linewidth=2.2, color=C_GF,
            label="Greedy-Refine", zorder=3)

    ax.set_xlabel("Merge Size (number of HotpotQA instances combined)",
                  fontsize=11, color="#2E3A59")
    ax.set_ylabel("Average Runtime (seconds)", fontsize=11, color="#2E3A59")
    ax.set_title(f"Runtime vs. Merge Size  —  Budget = {budget:,} tokens",
                 fontsize=13, fontweight="bold", color="#1A234E", pad=12)

    ax.set_xticks(merge_sizes)
    ax.tick_params(colors="#2E3A59", labelsize=10)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#CBD5E1")
    ax.spines["bottom"].set_color("#CBD5E1")
    ax.grid(axis="y", color="#CBD5E1", linewidth=0.8, linestyle="--", alpha=0.7)

    leg = ax.legend(fontsize=10, framealpha=0.9, edgecolor="#CBD5E1",
                    loc="upper left")
    for text in leg.get_texts():
        text.set_color("#2E3A59")

    # annotate peak DP value
    peak_dp  = max(dp_rt)
    peak_idx = dp_rt.index(peak_dp)
    ax.annotate(f"{peak_dp:.3f}s",
                xy=(merge_sizes[peak_idx], peak_dp),
                xytext=(merge_sizes[peak_idx] - 4, peak_dp + 0.05 * max(dp_rt)),
                fontsize=9, color=C_DP,
                arrowprops=dict(arrowstyle="-", color=C_DP, lw=0.8))

    plt.tight_layout(pad=1.2)
    plt.savefig(out_path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    print(f"  Chart saved: {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_c = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_c)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        elem = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            elem.set(qn(f"w:{k}"), v)
        tcBorders.append(elem)
    tcPr.append(tcBorders)


def add_paragraph(doc, text="", bold=False, italic=False,
                  size=None, color=None, align=None,
                  space_before=None, space_after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p


def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        if color:
            run.font.color.rgb = color
    return p


def callout_box(doc, text, label="Key Finding"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell  = table.cell(0, 0)
    set_cell_bg(cell, LIGHT)

    gold_hex = f"{GOLD[0]:02X}{GOLD[1]:02X}{GOLD[2]:02X}"
    teal_hex = f"{TEAL[0]:02X}{TEAL[1]:02X}{TEAL[2]:02X}"
    set_cell_border(cell,
        top    = {"sz": "8",  "val": "single", "color": gold_hex},
        bottom = {"sz": "8",  "val": "single", "color": gold_hex},
        left   = {"sz": "24", "val": "single", "color": teal_hex},
        right  = {"sz": "8",  "val": "single", "color": gold_hex},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(6)
    lr = p.add_run(f"▶  {label}:  ")
    lr.bold            = True
    lr.font.color.rgb  = TEAL
    lr.font.size       = Pt(10)
    br = p.add_run(text)
    br.font.size       = Pt(10)
    br.font.color.rgb  = SLATE
    doc.add_paragraph()


def styled_table(doc, headers, rows, caption=None):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], TH_BG)
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold           = True
                run.font.color.rgb = WHITE
                run.font.size      = Pt(9)

    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg    = ALT_ROW if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val)
            set_cell_bg(cells[c_idx], bg)
            for para in cells[c_idx].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)

    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.italic         = True
            run.font.size      = Pt(9)
            run.font.color.rgb = SLATE
        cp.paragraph_format.space_before = Pt(4)
    doc.add_paragraph()


def divider(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    teal_hex = f"{TEAL[0]:02X}{TEAL[1]:02X}{TEAL[2]:02X}"
    bot.set(qn("w:color"), teal_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


def embed_chart(doc, img_path, fig_num, caption_text):
    """Embed a PNG chart with a centred caption below it."""
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after  = Pt(2)
    run = p_img.add_run()
    run.add_picture(img_path, width=Inches(5.8))

    cap = doc.add_paragraph(f"Figure {fig_num}.  {caption_text}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic         = True
        run.font.size      = Pt(9)
        run.font.color.rgb = SLATE
    cap.paragraph_format.space_after = Pt(10)


# ─────────────────────────────────────────────────────────────────────────────
# Build document
# ─────────────────────────────────────────────────────────────────────────────

def build_report():
    # 1. Load CSV data
    print("Loading CSV data...")
    dp_rows = read_csv(CSV_DP)
    gr_rows = read_csv(CSV_GR)
    gf_rows = read_csv(CSV_GF)

    # 2. Generate charts
    print("Generating charts...")
    chart_paths = {}
    for budget in [2000, 4000, 8000]:
        path = os.path.join(REPORT, f"runtime_chart_{budget}.png")
        make_runtime_chart(dp_rows, gr_rows, gf_rows, budget, path)
        chart_paths[budget] = path

    # 3. Build DOCX
    print("Building DOCX...")
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── TITLE BLOCK ──
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(10)
    tp.paragraph_format.space_after  = Pt(4)
    tr = tp.add_run("LLM Context Compression\nas a 0/1 Knapsack Problem")
    tr.bold = True; tr.font.size = Pt(22); tr.font.color.rgb = NAVY; tr.font.name = "Calibri"

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(2)
    sr = sp.add_run("Algorithms Project — Final Experiment Report")
    sr.italic = True; sr.font.size = Pt(13); sr.font.color.rgb = TEAL; sr.font.name = "Calibri"

    dp2 = doc.add_paragraph()
    dp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp2.paragraph_format.space_after = Pt(16)
    dr2 = dp2.add_run("March 2026")
    dr2.font.size = Pt(10); dr2.font.color.rgb = SLATE

    divider(doc)

    # ── ABSTRACT ──
    heading(doc, "Abstract", level=1, color=NAVY)
    add_paragraph(doc,
        "This report presents a study of LLM context compression modeled as a 0/1 knapsack "
        "optimization problem. Given a question and a pool of text chunks derived from multiple "
        "Wikipedia passages, the goal is to select the highest-utility subset of chunks that fits "
        "within a fixed token budget. We evaluate three algorithms — exact dynamic programming "
        "(Exact-DP), greedy-by-ratio (Greedy-Ratio), and greedy-with-local-refinement "
        "(Greedy-Refine) — across 100 merged HotpotQA instances spanning five merge sizes and "
        "three token budgets. Our findings show that Greedy-Ratio achieves near-identical utility "
        "to Exact-DP while running up to three orders of magnitude faster, making it a compelling "
        "practical choice for real-time context management in large language model pipelines.",
        size=11, space_after=6)
    divider(doc)

    # ── INTRODUCTION ──
    heading(doc, "1  Introduction", level=1, color=NAVY)
    add_paragraph(doc,
        "Large language models (LLMs) operate under strict token-budget constraints imposed by "
        "finite context windows. When a retrieval-augmented generation (RAG) pipeline retrieves "
        "more text than fits in the context window, some form of context compression is required. "
        "A principled approach is to formulate this selection as a combinatorial optimization "
        "problem: assign each candidate text chunk a token cost and a relevance utility score, "
        "then select the feasible subset that maximizes total utility without exceeding the "
        "token budget.",
        size=11, space_after=6)
    add_paragraph(doc,
        "This framing maps directly onto the classical 0/1 Knapsack Problem. The knapsack "
        "capacity is the token budget; each item is a sentence-level chunk with a weight "
        "(token count) and a value (lexical relevance score). This study implements and "
        "benchmarks three solution strategies of increasing practical interest: an exact DP "
        "baseline, a fast greedy heuristic, and a hybrid refinement approach. The experiment "
        "is conducted on a curated dataset derived from HotpotQA to reflect realistic "
        "multi-document retrieval scenarios.",
        size=11, space_after=6)
    divider(doc)

    # ── PROBLEM FORMULATION ──
    heading(doc, "2  Problem Formulation", level=1, color=NAVY)
    add_paragraph(doc,
        "Let C = {c₁, c₂, …, cₙ} denote the set of candidate text chunks obtained by "
        "splitting retrieved documents at sentence boundaries. Each chunk cᵢ is characterized by:",
        size=11, space_after=4)
    for item in [
        "wᵢ ∈ ℤ⁺  — token count of the chunk (item weight)",
        "vᵢ ∈ ℝ≥0  — lexical relevance score against the query (item value)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item).font.size = Pt(11)

    p_obj = doc.add_paragraph()
    p_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_obj.paragraph_format.space_before = Pt(8)
    p_obj.paragraph_format.space_after  = Pt(8)
    r1 = p_obj.add_run("maximize  Σᵢ xᵢ · vᵢ     subject to  Σᵢ xᵢ · wᵢ ≤ B,    xᵢ ∈ {0, 1}")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = NAVY

    add_paragraph(doc,
        "where xᵢ = 1 indicates that chunk cᵢ is included in the compressed context. "
        "This is precisely the 0/1 Knapsack Problem, which is NP-hard in general but "
        "admits a pseudo-polynomial exact solution via dynamic programming in O(n · B) time and space.",
        size=11, space_after=6)
    divider(doc)

    # ── DATASET CONSTRUCTION ──
    heading(doc, "3  Dataset Construction", level=1, color=NAVY)
    add_paragraph(doc,
        "Experiments are conducted on instances derived from the HotpotQA distractor validation "
        "split, a multi-hop question-answering dataset where each question requires reasoning "
        "over two Wikipedia passages while eight distractor passages are included as noise.",
        size=11, space_after=6)
    heading(doc, "3.1  Merged Instances", level=2, color=SLATE)
    add_paragraph(doc,
        "To simulate increasingly large retrieval pools, we constructed merged instances by "
        "combining multiple HotpotQA examples while retaining a single target question-answer "
        "pair. Each merged instance therefore contains many more passages than a single HotpotQA "
        "item — most of which are distractors — creating a challenging and realistic "
        "compression scenario.",
        size=11, space_after=6)
    styled_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ["Merge sizes", "10, 20, 30, 40, 50"],
            ["Samples per merge size", "20"],
            ["Total merged instances", "100"],
            ["Chunking granularity", "Sentence-level"],
            ["Utility scoring method", "Lexical relevance (TF-IDF-inspired overlap)"],
            ["Token budgets evaluated", "2,000 · 4,000 · 8,000 tokens"],
            ["Cost scaling (dp_cost_scale)", "1 (raw token costs — all methods identical)"],
        ],
        caption="Table 1  —  Dataset and experimental configuration summary."
    )
    divider(doc)

    # ── EXPERIMENTAL SETUP ──
    heading(doc, "4  Experimental Setup", level=1, color=NAVY)
    add_paragraph(doc,
        "Each of the 100 merged instances was solved by all three algorithms at every combination "
        "of token budget and merge size, yielding 15 distinct (merge_size, budget) conditions "
        "(5 × 3). Runtime was measured using Python's high-resolution perf_counter and averaged "
        "across the 20 samples in each condition. All methods received identical chunk sets and "
        "utility scores to ensure a fair comparison. Brute-force enumeration was excluded from "
        "the full experiment because its O(2ⁿ) complexity renders it computationally infeasible "
        "on instances with hundreds of chunks.",
        size=11, space_after=6)
    divider(doc)

    # ── ALGORITHMS COMPARED ──
    heading(doc, "5  Algorithms Compared", level=1, color=NAVY)
    heading(doc, "5.1  Exact Dynamic Programming (Exact-DP)", level=2, color=SLATE)
    add_paragraph(doc,
        "Exact-DP solves the 0/1 Knapsack Problem optimally using a standard bottom-up DP table "
        "of size (n+1) × (B+1). It guarantees the globally optimal selection but incurs O(n · B) "
        "time and space complexity. As n (chunks) and B (budget) grow, this becomes the "
        "performance bottleneck.", size=11, space_after=6)
    heading(doc, "5.2  Greedy by Utility-to-Cost Ratio (Greedy-Ratio)", level=2, color=SLATE)
    add_paragraph(doc,
        "Greedy-Ratio sorts chunks in descending order of value per token (vᵢ / wᵢ) and "
        "greedily includes each chunk while remaining capacity allows. This runs in O(n log n) "
        "time — dominated by sorting — and requires O(n) space. The approximation is "
        "well-studied; it can miss the optimum but performs very well in practice when items "
        "are numerous and values are spread across a continuous range.", size=11, space_after=6)
    heading(doc, "5.3  Greedy with Local Refinement (Greedy-Refine)", level=2, color=SLATE)
    add_paragraph(doc,
        "Greedy-Refine begins with the Greedy-Ratio solution and then attempts to improve it "
        "by swapping selected chunks for unselected ones when a swap increases total utility "
        "within budget. This adds a local-search pass on top of the greedy initialisation. "
        "In theory it can recover some of the optimality gap; in practice the benefit depends "
        "on the instance structure and refinement depth.", size=11, space_after=6)
    divider(doc)

    # ── EVALUATION METRICS ──
    heading(doc, "6  Evaluation Metrics", level=1, color=NAVY)
    styled_table(doc,
        headers=["Metric", "Description"],
        rows=[
            ["Avg. Runtime (s)",            "Mean wall-clock time per instance"],
            ["Avg. Selected Utility",        "Mean total utility of the selected chunk set"],
            ["Avg. Support Recall",          "Fraction of gold-support sentences covered"],
            ["Avg. Exact Support Coverage",  "Proportion of instances with all supports selected"],
            ["Avg. Budget Utilization",      "Fraction of the token budget consumed"],
            ["Avg. Compression Ratio",       "Ratio of selected tokens to total available tokens"],
            ["Avg. Selected Chunks",         "Mean number of chunks chosen per instance"],
        ],
        caption="Table 2  —  Evaluation metrics used in the full experiment."
    )
    divider(doc)

    # ── RESULTS ──
    heading(doc, "7  Results", level=1, color=NAVY)
    add_paragraph(doc,
        "Tables 3–5 present the aggregated results for all three algorithms across all "
        "15 (merge_size, budget) conditions. Numbers are averaged over 20 samples.",
        size=11, space_after=8)

    heading(doc, "7.1  Exact-DP Results", level=2, color=SLATE)
    styled_table(doc,
        headers=["Merge", "Budget", "Runtime (s)", "Utility", "Support Recall",
                 "Exact Cov.", "Budget Util.", "Comp. Ratio"],
        rows=[
            [10,"2,000","0.121","18.99","0.875","0.65","1.000","0.182"],
            [10,"4,000","0.270","26.75","0.885","0.70","1.000","0.363"],
            [10,"8,000","0.564","35.47","0.920","0.75","0.995","0.722"],
            [20,"2,000","0.242","21.98","0.812","0.50","1.000","0.091"],
            [20,"4,000","0.516","31.42","0.922","0.75","1.000","0.181"],
            [20,"8,000","1.073","44.10","0.947","0.80","1.000","0.362"],
            [30,"2,000","0.333","23.10","0.838","0.65","1.000","0.061"],
            [30,"4,000","0.709","33.15","0.879","0.70","1.000","0.121"],
            [30,"8,000","1.506","46.91","0.904","0.75","0.999","0.242"],
            [40,"2,000","0.475","21.40","0.783","0.55","1.000","0.045"],
            [40,"4,000","1.015","31.24","0.853","0.65","1.000","0.091"],
            [40,"8,000","2.067","45.06","0.878","0.70","1.000","0.182"],
            [50,"2,000","0.503","22.81","0.782","0.50","1.000","0.037"],
            [50,"4,000","1.049","33.75","0.832","0.60","1.000","0.074"],
            [50,"8,000","2.136","48.80","0.857","0.65","0.989","0.146"],
        ],
        caption="Table 3  —  Exact-DP results across all conditions (20 samples each)."
    )

    heading(doc, "7.2  Greedy-Ratio Results", level=2, color=SLATE)
    styled_table(doc,
        headers=["Merge","Budget","Runtime (s)","Utility","Support Recall",
                 "Exact Cov.","Budget Util.","Comp. Ratio"],
        rows=[
            [10,"2,000","0.0003","18.98","0.875","0.65","1.000","0.182"],
            [10,"4,000","0.0003","26.74","0.885","0.70","1.000","0.363"],
            [10,"8,000","0.0003","35.46","0.920","0.75","1.000","0.726"],
            [20,"2,000","0.0006","21.96","0.828","0.55","1.000","0.091"],
            [20,"4,000","0.0005","31.42","0.922","0.75","1.000","0.181"],
            [20,"8,000","0.0006","44.10","0.947","0.80","1.000","0.362"],
            [30,"2,000","0.0011","23.09","0.838","0.65","1.000","0.061"],
            [30,"4,000","0.0009","33.14","0.879","0.70","1.000","0.121"],
            [30,"8,000","0.0009","46.90","0.904","0.75","1.000","0.243"],
            [40,"2,000","0.0015","21.39","0.783","0.55","1.000","0.045"],
            [40,"4,000","0.0012","31.23","0.853","0.65","1.000","0.091"],
            [40,"8,000","0.0012","45.05","0.878","0.70","1.000","0.182"],
            [50,"2,000","0.0020","22.80","0.782","0.50","1.000","0.037"],
            [50,"4,000","0.0016","33.73","0.832","0.60","1.000","0.074"],
            [50,"8,000","0.0017","48.79","0.873","0.70","1.000","0.147"],
        ],
        caption="Table 4  —  Greedy-Ratio results across all conditions (20 samples each)."
    )

    heading(doc, "7.3  Greedy-Refine Results", level=2, color=SLATE)
    styled_table(doc,
        headers=["Merge","Budget","Runtime (s)","Utility","Support Recall",
                 "Exact Cov.","Budget Util.","Comp. Ratio"],
        rows=[
            [10,"2,000","0.0059","18.98","0.875","0.65","1.000","0.182"],
            [10,"4,000","0.0089","26.74","0.885","0.70","1.000","0.363"],
            [10,"8,000","0.0102","35.46","0.920","0.75","1.000","0.726"],
            [20,"2,000","0.0060","21.96","0.828","0.55","1.000","0.091"],
            [20,"4,000","0.0104","31.42","0.922","0.75","1.000","0.181"],
            [20,"8,000","0.0187","44.10","0.947","0.80","1.000","0.362"],
            [30,"2,000","0.0061","23.09","0.838","0.65","1.000","0.061"],
            [30,"4,000","0.0102","33.14","0.879","0.70","1.000","0.121"],
            [30,"8,000","0.0181","46.90","0.904","0.75","1.000","0.243"],
            [40,"2,000","0.0068","21.39","0.783","0.55","1.000","0.045"],
            [40,"4,000","0.0106","31.23","0.853","0.65","1.000","0.091"],
            [40,"8,000","0.0179","45.05","0.878","0.70","1.000","0.182"],
            [50,"2,000","0.0075","22.80","0.782","0.50","1.000","0.037"],
            [50,"4,000","0.0107","33.73","0.832","0.60","1.000","0.074"],
            [50,"8,000","0.0187","48.79","0.873","0.70","1.000","0.147"],
        ],
        caption="Table 5  —  Greedy-Refine results across all conditions (20 samples each)."
    )

    callout_box(doc,
        "Across all 15 conditions, Greedy-Ratio and Greedy-Refine achieved virtually identical "
        "utility, support recall, and exact coverage to Exact-DP — at a fraction of the "
        "computational cost.",
        label="Key Finding")
    divider(doc)

    # ── RUNTIME ANALYSIS ──
    heading(doc, "8  Runtime Analysis", level=1, color=NAVY)
    add_paragraph(doc,
        "Runtime is the primary axis of differentiation across the three algorithms. "
        "Figures 1–3 plot average runtime as a function of merge size for each of the "
        "three token budgets.",
        size=11, space_after=6)

    fig_captions = {
        2000: "Average runtime (seconds) vs. merge size for a token budget of 2,000 tokens. "
              "Exact-DP already exhibits noticeable scaling; the greedy methods remain near-flat.",
        4000: "Average runtime vs. merge size for a token budget of 4,000 tokens. "
              "The Exact-DP curve steepens further; greedy methods are still essentially instant.",
        8000: "Average runtime vs. merge size for a token budget of 8,000 tokens. "
              "At the hardest setting Exact-DP reaches 2.14 s at merge_size = 50, "
              "while Greedy-Ratio stays at 1.7 ms.",
    }
    for fig_num, budget in enumerate([2000, 4000, 8000], start=1):
        embed_chart(doc, chart_paths[budget], fig_num, fig_captions[budget])

    heading(doc, "8.1  Runtime Summary at the Hardest Setting", level=2, color=SLATE)
    styled_table(doc,
        headers=["Algorithm", "Avg Runtime (s)  [merge=50, budget=8,000]", "Speedup vs. Exact-DP"],
        rows=[
            ["Exact-DP",      "2.1355",  "1×  (baseline)"],
            ["Greedy-Ratio",  "0.0017",  "≈ 1,275×  faster"],
            ["Greedy-Refine", "0.0187",  "≈  114×  faster"],
        ],
        caption="Table 6  —  Runtime comparison at the most demanding setting "
                "(merge_size = 50, budget = 8,000 tokens)."
    )

    callout_box(doc,
        "At merge_size = 50 and budget = 8,000, Greedy-Ratio was approximately 1,275× faster "
        "than Exact-DP, and Greedy-Refine was roughly 114× faster — while both greedy methods "
        "maintained near-identical quality.",
        label="Key Finding")

    heading(doc, "8.2  Why Exact-DP Scales Poorly", level=2, color=SLATE)
    add_paragraph(doc,
        "Exact-DP constructs a DP table of size (n + 1) × (B + 1), where n is the number of "
        "chunks and B is the token budget. As merge size grows, n grows proportionally (more "
        "passages yield more sentence chunks). Simultaneously, larger budgets directly increase B. "
        "The combined effect is quadratic growth in the table size, which explains the steep "
        "runtime curves in Figures 1–3: doubling the merge size roughly doubles n and the "
        "number of DP iterations, while moving from budget 2,000 to 8,000 multiplies the inner "
        "dimension by four. The worst observed mean runtime — 2.14 s at merge_size = 50, "
        "budget = 8,000 — is already beyond real-time thresholds for interactive LLM systems.",
        size=11, space_after=6)

    heading(doc, "8.3  Why Greedy-Ratio Is Extremely Fast", level=2, color=SLATE)
    add_paragraph(doc,
        "Greedy-Ratio requires only a single sort of the chunks (O(n log n)) followed by a "
        "single linear scan. There is no multi-dimensional table to fill; the budget is consumed "
        "incrementally in one pass. Consequently, runtime scales only with the number of chunks, "
        "and that scaling is very gentle. Even at merge_size = 50, the average runtime is "
        "1.7 ms — essentially negligible in any practical pipeline.",
        size=11, space_after=6)

    heading(doc, "8.4  Why Greedy-Refine Did Not Materially Outperform Greedy-Ratio", level=2, color=SLATE)
    add_paragraph(doc,
        "Greedy-Refine's refinement pass attempts pairwise swaps to escape the greedy solution, "
        "but its effectiveness depends on whether improving swaps actually exist. In this "
        "experiment, the greedy solution was apparently already close to optimal for the majority "
        "of instances: the lexical utility scores are diverse enough that the ratio-ordered "
        "greedy selection captures the highest-value chunks without leaving many profitable swap "
        "opportunities. Because no improvement was found, the refinement pass added latency "
        "(≈ 10–19 ms per instance) without any corresponding gain in utility or recall, yielding "
        "results virtually identical to Greedy-Ratio.",
        size=11, space_after=6)
    divider(doc)

    # ── DISCUSSION ──
    heading(doc, "9  Discussion", level=1, color=NAVY)
    add_paragraph(doc,
        "The results tell a consistent story across all experimental conditions: the greedy "
        "approach achieves near-optimal quality at a dramatically lower computational cost. "
        "This reflects a known property of the knapsack problem in practice: when utility "
        "values are distributed broadly and item sizes are small relative to the budget, the "
        "fractional relaxation is tight and the greedy-ratio heuristic closely approximates "
        "the integral optimum. The retrieved text chunks in this dataset appear to satisfy "
        "both conditions.",
        size=11, space_after=6)
    add_paragraph(doc,
        "The practical implication is significant. A real-time RAG pipeline cannot afford "
        "2 seconds of compression latency per query. The greedy approach — at 1–2 ms — is "
        "effectively free relative to retrieval and generation costs. The optimality guarantees "
        "of Exact-DP, while theoretically appealing, provide no measurable quality benefit in "
        "this setting and carry a prohibitive computational cost at scale.",
        size=11, space_after=6)
    add_paragraph(doc,
        "Support recall consistently improved with larger token budgets, as expected: more "
        "budget allows more support sentences to be included. Exact support coverage was highest "
        "at merge_size = 10 and budget = 8,000, where the compressed context was large enough "
        "to encompass all gold-support material. As merge size grew, the distractor pool expanded, "
        "making full coverage harder even at large budgets.",
        size=11, space_after=6)
    divider(doc)

    # ── LIMITATIONS ──
    heading(doc, "10  Limitations", level=1, color=NAVY)
    for title, body in [
        ("Lexical utility only.",
         "The utility function is based on lexical overlap. Semantic similarity (e.g., dense "
         "embeddings) might better capture relevance for paraphrased or multi-hop reasoning."),
        ("Shallow refinement.",
         "The local-search refinement may be too shallow. Deeper strategies — iterated local "
         "search, simulated annealing, or branch-and-bound pruning — could potentially close "
         "the gap with Exact-DP on harder instances."),
        ("Controlled dataset.",
         "The merged HotpotQA instances are constructed rather than drawn from a production "
         "retrieval system. Real-world chunk distributions may differ in ways that affect "
         "algorithm ranking."),
        ("No brute-force comparison.",
         "Brute force was excluded due to exponential complexity. A small-scale comparison "
         "confirming the optimality of Exact-DP outputs was not part of this study."),
        ("No downstream LLM evaluation.",
         "Metrics are information-theoretic (recall, utility). Whether improved selection "
         "translates to better LLM answer accuracy remains an open empirical question."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(title + "  "); r_t.bold = True; r_t.font.size = Pt(11)
        r_b = p.add_run(body); r_b.font.size = Pt(11)

    doc.add_paragraph()
    divider(doc)

    # ── CONCLUSION ──
    heading(doc, "11  Conclusion", level=1, color=NAVY)
    add_paragraph(doc,
        "This project demonstrates that context compression for LLMs can be rigorously framed "
        "as a 0/1 Knapsack Problem, enabling the direct application of classical algorithm "
        "design theory to a pressing practical challenge in natural language processing.",
        size=11, space_after=6)
    add_paragraph(doc,
        "The central finding is that algorithmic efficiency and solution quality are not at odds "
        "in this domain. Greedy-Ratio, a simple O(n log n) heuristic, matches the quality of "
        "an exact pseudo-polynomial DP solver across every measured dimension — utility, support "
        "recall, budget utilization — while running more than three orders of magnitude faster "
        "at the largest tested scale. This gap will only widen as context windows grow and "
        "retrieval pools expand.",
        size=11, space_after=6)
    add_paragraph(doc,
        "The takeaway is both algorithmic and practical: when designing compression pipelines "
        "for production LLM systems, a well-chosen greedy heuristic is not a shortcut — it is "
        "the right tool. The pseudo-polynomial cost of exact methods buys no measurable NLP "
        "benefit in this setting, confirming that the theoretical worst-case approximation gap "
        "of greedy knapsack algorithms rarely materializes on structured natural-language "
        "instances. Future work should explore richer utility functions, adaptive budgeting, "
        "and the impact of context quality on downstream LLM generation accuracy.",
        size=11, space_after=8)

    callout_box(doc,
        "Framing LLM context compression as a knapsack problem reveals that the greedy "
        "heuristic is not merely 'good enough' — it is essentially optimal at a fraction of "
        "the cost. Algorithmic complexity theory predicts exactly this, and the empirical data "
        "confirms it.",
        label="Concluding Takeaway")

    doc.save(OUT)
    print(f"\nReport saved to: {OUT}")


if __name__ == "__main__":
    build_report()
