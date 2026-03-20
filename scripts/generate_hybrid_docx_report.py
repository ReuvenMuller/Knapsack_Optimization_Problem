from __future__ import annotations

import csv
import os
from datetime import datetime, timezone
from statistics import mean

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT_DIR = os.path.join(ROOT_DIR, "report")
FIG_DIR = os.path.join(REPORT_DIR, "figures")


def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def read_csv(path: str) -> list[dict[str, str]]:
    with open(path, "r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def to_float(row: dict[str, str], key: str) -> float:
    value = row.get(key, "")
    return float(value) if value not in ("", None) else 0.0


def group_by_budget(rows: list[dict[str, str]], metric: str) -> dict[int, float]:
    budgets = sorted({int(row["budget_tokens"]) for row in rows})
    return {
        budget: mean(to_float(row, metric) for row in rows if int(row["budget_tokens"]) == budget)
        for budget in budgets
    }


def overall_metric(rows: list[dict[str, str]], metric: str) -> float:
    return mean(to_float(row, metric) for row in rows)


def plot_runtime_all_algorithms(
    budget: int,
    exact_rows: list[dict[str, str]],
    greedy_rows: list[dict[str, str]],
    refine_rows: list[dict[str, str]],
    output_path: str,
) -> None:
    fig, ax = plt.subplots(figsize=(9, 5.2))
    for label, rows, color in [
        ("exact_dp", exact_rows, "#0B5D8C"),
        ("greedy_ratio", greedy_rows, "#2C9A5F"),
        ("greedy_refine", refine_rows, "#C06C2B"),
    ]:
        xs = [int(row["merge_size"]) for row in rows if int(row["budget_tokens"]) == budget]
        ys = [to_float(row, "avg_runtime_sec") for row in rows if int(row["budget_tokens"]) == budget]
        ax.plot(xs, ys, marker="o", linewidth=2.5, label=label, color=color)
    ax.set_title(f"Hybrid Utility: Runtime vs Merge Size (Budget = {budget})")
    ax.set_xlabel("Merge Size")
    ax.set_ylabel("Average Runtime (seconds)")
    ax.grid(True, alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def plot_runtime_greedy_only(
    budget: int,
    greedy_rows: list[dict[str, str]],
    refine_rows: list[dict[str, str]],
    output_path: str,
) -> None:
    fig, ax = plt.subplots(figsize=(9, 5.2))
    for label, rows, color in [
        ("greedy_ratio", greedy_rows, "#2C9A5F"),
        ("greedy_refine", refine_rows, "#C06C2B"),
    ]:
        xs = [int(row["merge_size"]) for row in rows if int(row["budget_tokens"]) == budget]
        ys = [to_float(row, "avg_runtime_sec") for row in rows if int(row["budget_tokens"]) == budget]
        ax.plot(xs, ys, marker="o", linewidth=2.5, label=label, color=color)
    ax.set_title(f"Hybrid Utility: Greedy Runtime Comparison (Budget = {budget})")
    ax.set_xlabel("Merge Size")
    ax.set_ylabel("Average Runtime (seconds)")
    ax.grid(True, alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def plot_utility_comparison(
    lexical_rows_by_algo: dict[str, list[dict[str, str]]],
    hybrid_rows_by_algo: dict[str, list[dict[str, str]]],
    metric: str,
    title: str,
    ylabel: str,
    output_path: str,
) -> None:
    algos = ["exact_dp", "greedy_ratio", "greedy_refine"]
    lexical_vals = [overall_metric(lexical_rows_by_algo[algo], metric) for algo in algos]
    hybrid_vals = [overall_metric(hybrid_rows_by_algo[algo], metric) for algo in algos]

    fig, ax = plt.subplots(figsize=(8.5, 5.2))
    positions = range(len(algos))
    width = 0.35
    ax.bar([x - width / 2 for x in positions], lexical_vals, width=width, label="lexical", color="#7A8DA6")
    ax.bar([x + width / 2 for x in positions], hybrid_vals, width=width, label="hybrid", color="#2C9A5F")
    ax.set_xticks(list(positions))
    ax.set_xticklabels(algos)
    ax.set_title(title)
    ax.set_ylabel(ylabel)
    ax.grid(True, axis="y", alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text


def add_figure(doc: Document, path: str, caption: str) -> None:
    doc.add_picture(path, width=Inches(6.8))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    ensure_dir(REPORT_DIR)
    ensure_dir(FIG_DIR)

    lexical = {
        "exact_dp": read_csv(os.path.join(ROOT_DIR, "results", "exact_dp", "fresh_full_exact_dp", "summary.csv")),
        "greedy_ratio": read_csv(os.path.join(ROOT_DIR, "results", "greedy_ratio", "fresh_full_greedy_ratio", "summary.csv")),
        "greedy_refine": read_csv(os.path.join(ROOT_DIR, "results", "greedy_refine", "fresh_full_greedy_refine", "summary.csv")),
    }
    hybrid = {
        "exact_dp": read_csv(os.path.join(ROOT_DIR, "results", "exact_dp", "hybrid_full_exact_dp", "summary.csv")),
        "greedy_ratio": read_csv(os.path.join(ROOT_DIR, "results", "greedy_ratio", "hybrid_full_greedy_ratio", "summary.csv")),
        "greedy_refine": read_csv(os.path.join(ROOT_DIR, "results", "greedy_refine", "hybrid_full_greedy_refine", "summary.csv")),
    }

    figures = {
        "runtime_2000": os.path.join(FIG_DIR, "hybrid_runtime_budget_2000.png"),
        "runtime_4000": os.path.join(FIG_DIR, "hybrid_runtime_budget_4000.png"),
        "runtime_8000": os.path.join(FIG_DIR, "hybrid_runtime_budget_8000.png"),
        "greedy_only_8000": os.path.join(FIG_DIR, "hybrid_greedy_only_budget_8000.png"),
        "coverage_compare": os.path.join(FIG_DIR, "lexical_vs_hybrid_exact_coverage.png"),
        "recall_compare": os.path.join(FIG_DIR, "lexical_vs_hybrid_support_recall.png"),
    }

    plot_runtime_all_algorithms(2000, hybrid["exact_dp"], hybrid["greedy_ratio"], hybrid["greedy_refine"], figures["runtime_2000"])
    plot_runtime_all_algorithms(4000, hybrid["exact_dp"], hybrid["greedy_ratio"], hybrid["greedy_refine"], figures["runtime_4000"])
    plot_runtime_all_algorithms(8000, hybrid["exact_dp"], hybrid["greedy_ratio"], hybrid["greedy_refine"], figures["runtime_8000"])
    plot_runtime_greedy_only(8000, hybrid["greedy_ratio"], hybrid["greedy_refine"], figures["greedy_only_8000"])
    plot_utility_comparison(
        lexical,
        hybrid,
        "avg_exact_support_coverage",
        "Lexical vs Hybrid: Average Exact Support Coverage",
        "Average Exact Support Coverage",
        figures["coverage_compare"],
    )
    plot_utility_comparison(
        lexical,
        hybrid,
        "avg_support_recall",
        "Lexical vs Hybrid: Average Support Recall",
        "Average Support Recall",
        figures["recall_compare"],
    )

    # Key comparisons.
    hardest_exact = next(row for row in hybrid["exact_dp"] if int(row["merge_size"]) == 50 and int(row["budget_tokens"]) == 8000)
    hardest_greedy = next(row for row in hybrid["greedy_ratio"] if int(row["merge_size"]) == 50 and int(row["budget_tokens"]) == 8000)
    hardest_refine = next(row for row in hybrid["greedy_refine"] if int(row["merge_size"]) == 50 and int(row["budget_tokens"]) == 8000)
    speedup_greedy = to_float(hardest_exact, "avg_runtime_sec") / to_float(hardest_greedy, "avg_runtime_sec")
    speedup_refine = to_float(hardest_exact, "avg_runtime_sec") / to_float(hardest_refine, "avg_runtime_sec")

    overall_rows = []
    for algo in ["exact_dp", "greedy_ratio", "greedy_refine"]:
        overall_rows.append([
            algo,
            f"{overall_metric(lexical[algo], 'avg_support_recall'):.3f}",
            f"{overall_metric(hybrid[algo], 'avg_support_recall'):.3f}",
            f"{overall_metric(lexical[algo], 'avg_exact_support_coverage'):.3f}",
            f"{overall_metric(hybrid[algo], 'avg_exact_support_coverage'):.3f}",
        ])

    md_path = os.path.join(REPORT_DIR, "hybrid_comparison_report.md")
    md_lines: list[str] = []
    md_lines.append("# Hybrid Utility Report")
    md_lines.append("")
    md_lines.append(f"Generated on {datetime.now(timezone.utc).isoformat()}")
    md_lines.append("")
    md_lines.append("## Summary")
    md_lines.append("")
    md_lines.append("- Main focus: hybrid utility (`0.7 * semantic + 0.3 * lexical`) using `sentence-transformers/all-MiniLM-L6-v2`.")
    md_lines.append("- Hybrid results are compared against the earlier lexical-only runs.")
    md_lines.append(
        f"- At `merge_size = 50`, `budget = 8000`, hybrid `exact_dp` averaged {to_float(hardest_exact, 'avg_runtime_sec'):.4f}s, "
        f"`greedy_ratio` averaged {to_float(hardest_greedy, 'avg_runtime_sec'):.4f}s, and `greedy_refine` averaged {to_float(hardest_refine, 'avg_runtime_sec'):.4f}s."
    )
    md_lines.append(
        f"- In that hardest hybrid setting, `greedy_ratio` was about {speedup_greedy:.0f}x faster than `exact_dp`, and `greedy_refine` was about {speedup_refine:.0f}x faster."
    )
    md_lines.append("- Hybrid generally improved support recall and exact support coverage relative to lexical, especially at medium and high budgets.")
    md_lines.append("- `greedy_ratio` and `greedy_refine` remained very close in quality, so the extra local-search cost did not translate into a large aggregate gain.")
    md_lines.append("")
    md_lines.append("## Overall Lexical vs Hybrid Comparison")
    md_lines.append("")
    md_lines.append("| Algorithm | Lexical Recall | Hybrid Recall | Lexical Exact Coverage | Hybrid Exact Coverage |")
    md_lines.append("|---|---:|---:|---:|---:|")
    for row in overall_rows:
        md_lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} |")
    md_lines.append("")
    md_lines.append("## Figures")
    md_lines.append("")
    for name in ["runtime_2000", "runtime_4000", "runtime_8000", "greedy_only_8000", "coverage_compare", "recall_compare"]:
        md_lines.append(f"![{name}]({figures[name]})")
        md_lines.append("")
    with open(md_path, "w", encoding="utf-8") as handle:
        handle.write("\n".join(md_lines))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Knapsack Context Compression\nHybrid Utility Report")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Full experiment summary with lexical vs hybrid comparison")
    run.italic = True
    run.font.size = Pt(11)

    add_heading(doc, "1. Overview", 1)
    add_bullet(doc, "Dataset: merged HotpotQA distractor validation set")
    add_bullet(doc, "100 merged instances across merge sizes 10, 20, 30, 40, and 50")
    add_bullet(doc, "Budgets: 2000, 4000, 8000")
    add_bullet(doc, "Algorithms: exact_dp, greedy_ratio, greedy_refine")
    add_bullet(doc, "Hybrid utility: 0.7 semantic + 0.3 lexical")
    add_bullet(doc, "Semantic model: sentence-transformers/all-MiniLM-L6-v2")

    add_heading(doc, "2. Key Findings", 1)
    add_bullet(
        doc,
        f"At merge size 50 and budget 8000, hybrid exact_dp averaged {to_float(hardest_exact, 'avg_runtime_sec'):.4f}s, "
        f"greedy_ratio averaged {to_float(hardest_greedy, 'avg_runtime_sec'):.4f}s, and greedy_refine averaged {to_float(hardest_refine, 'avg_runtime_sec'):.4f}s.",
    )
    add_bullet(doc, f"Greedy_ratio was about {speedup_greedy:.0f}x faster than exact_dp in that hardest setting.")
    add_bullet(doc, f"Greedy_refine was about {speedup_refine:.0f}x faster than exact_dp in that hardest setting.")
    add_bullet(doc, "Hybrid generally improved support recall and exact support coverage compared with lexical.")
    add_bullet(doc, "Greedy_refine remained close to greedy_ratio in aggregate quality, so the extra runtime yielded limited additional benefit.")

    add_heading(doc, "3. Lexical vs Hybrid Comparison", 1)
    add_table(
        doc,
        ["Algorithm", "Lexical Recall", "Hybrid Recall", "Lexical Exact Coverage", "Hybrid Exact Coverage"],
        overall_rows,
    )
    add_paragraph(
        doc,
        "Interpretation: hybrid scoring appears to align more closely with the supporting evidence than lexical-only scoring. "
        "The strongest gains show up in exact support coverage, which is a strict metric and therefore meaningful for this project.",
    )

    add_heading(doc, "4. Runtime Analysis", 1)
    add_paragraph(
        doc,
        "The runtime story remains clean. Exact dynamic programming scales the worst as merge size and budget increase. "
        "Greedy ratio remains extremely fast, while greedy refinement adds noticeable overhead but stays far below DP.",
    )
    add_figure(doc, figures["runtime_2000"], "Figure 1. Hybrid utility runtime vs merge size at budget 2000.")
    add_figure(doc, figures["runtime_4000"], "Figure 2. Hybrid utility runtime vs merge size at budget 4000.")
    add_figure(doc, figures["runtime_8000"], "Figure 3. Hybrid utility runtime vs merge size at budget 8000.")
    add_figure(doc, figures["greedy_only_8000"], "Figure 4. Greedy-only runtime comparison at budget 8000, shown separately so the difference is visible without DP dominating the scale.")

    add_heading(doc, "5. Utility Comparison Figures", 1)
    add_figure(doc, figures["coverage_compare"], "Figure 5. Average exact support coverage: lexical vs hybrid.")
    add_figure(doc, figures["recall_compare"], "Figure 6. Average support recall: lexical vs hybrid.")

    add_heading(doc, "6. Conclusion", 1)
    add_paragraph(
        doc,
        "The new hybrid utility results strengthen the project. They preserve the same overall runtime hierarchy while improving evidence-retention metrics over lexical-only scoring. "
        "For this implementation, greedy_ratio remains the best practical speed-quality tradeoff, while greedy_refine does not yet show a strong enough improvement to justify its extra runtime.",
    )

    docx_path = os.path.join(REPORT_DIR, "hybrid_comparison_report.docx")
    doc.save(docx_path)

    print(f"Markdown report written to: {md_path}")
    print(f"DOCX report written to: {docx_path}")
    for path in figures.values():
        print(f"Figure written to: {path}")


if __name__ == "__main__":
    main()
